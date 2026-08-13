import pandas as pd
import time
from langchain_chroma import Chroma
import json
from langchain_core.documents import Document
from docx import Document as DocxDocument
import os
from dotenv import load_dotenv
from langchain_google_genai import (
    GoogleGenerativeAIEmbeddings,
    ChatGoogleGenerativeAI
)
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_core.documents import Document
from langchain_core.prompts import (
    ChatPromptTemplate,
    MessagesPlaceholder
)
from langchain_core.messages import (
    HumanMessage,
    AIMessage,
    AnyMessage
)
# --- NEW: BM25 (sparse/keyword) retriever + ensemble fusion ---
from langchain_community.retrievers import BM25Retriever
try:
    # langchain <1.0
    from langchain.retrievers import EnsembleRetriever
except ImportError:
    # langchain >=1.0 split EnsembleRetriever out into langchain_classic
    from langchain_classic.retrievers.ensemble import EnsembleRetriever

from langgraph.graph.message import add_messages
from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import StateGraph, START, END
from typing_extensions import TypedDict, Annotated
load_dotenv()

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")


def load_vector_store():

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001",
        google_api_key=GOOGLE_API_KEY
    )

    persist_directory = "./chroma_db_new"

    # LOAD CHROMADB
    vector_store = Chroma(
        collection_name="medical_data",
        embedding_function=embeddings,
        persist_directory=persist_directory
    )
    # GET EXISTING IDS
    existing = vector_store.get()

    existing_ids = set(existing["ids"])

    documents = []
    ids = []
    print(sorted(existing_ids))
    
    # LOAD EXCEL
    from pathlib import Path
    import pandas as pd
    BASE_DIR = Path(__file__).resolve().parent.parent

    excel_file = BASE_DIR / "documents" / "Medical_list_with_specs.csv"

    df = pd.read_csv(excel_file)

    for index, row in df.iterrows():
        row_dict = row.to_dict()
        chunk_text = json.dumps(row_dict)
        doc_id = f"csv_{index}"

        documents.append(
            Document(
                page_content=chunk_text,
                metadata={
                    "source": "csv",
                    "id": doc_id
                }
            )
        )
        ids.append(doc_id)
    
    # FILTER NEW DOCS
    
    new_docs = []
    new_ids = []

    for doc, doc_id in zip(documents, ids): 

        if doc_id not in existing_ids:

            new_docs.append(doc)
            new_ids.append(doc_id)

 # EMBED IN BATCHES OF 100
    batch_size = 100

    if new_docs:

        for start in range(0, len(new_docs), batch_size):

            end = start + batch_size

            batch_docs = new_docs[start:end]
            batch_ids = new_ids[start:end]

            print(
                f"Embedding batch "
                f"{start // batch_size + 1}: "
                f"{len(batch_docs)} documents"
            )

            vector_store.add_documents(
                documents=batch_docs,
                ids=batch_ids
            )

            # Wait 1 minute before the next batch
            if end < len(new_docs):
                print("Waiting 60 seconds before next batch...")
                time.sleep(60)

    else:
        print("No new CSV documents to embed")

    return vector_store


def load_bm25_retriever(k: int = 7):
    """
    Build a BM25 (sparse / keyword-overlap) retriever over every document
    currently sitting in the Chroma collection.

    We pull straight from Chroma (via .get()) instead of re-reading the CSV
    so this stays in sync with the *actual* embedded corpus -- CSV rows,
    docx chunks (FAQ / Log & sign), and anything embedded in future --
    without needing a second source of truth.

    BM25 is a lexical/term-frequency ranker. It's included alongside the
    dense embedding retriever because it's much stronger at exact-match
    lookups (model numbers, SKUs, spec keywords like "1000 VA" or
    "MRI compatible") that embeddings can sometimes blur or miss.
    """
    raw = vector_store.get(include=["documents", "metadatas"])

    docs = [
        Document(page_content=content, metadata=metadata or {})
        for content, metadata in zip(raw["documents"], raw["metadatas"])
    ]

    if not docs:
        # BM25Retriever.from_documents() errors out on an empty corpus,
        # so fall back to a single empty doc rather than crashing at import.
        docs = [Document(page_content="", metadata={})]

    bm25_retriever = BM25Retriever.from_documents(docs)
    bm25_retriever.k = k

    return bm25_retriever


vector_store = load_vector_store()

# --- NEW: hybrid retrieval setup ---
# Dense (semantic) retriever, wrapping the existing Chroma vector store.
vector_retriever = vector_store.as_retriever(search_kwargs={"k": 7})

# Sparse (keyword / BM25) retriever, built from the same corpus.
bm25_retriever = load_bm25_retriever(k=7)

# Fuse both rankings via Reciprocal Rank Fusion. Weights favor the dense
# retriever slightly (0.6) while still letting BM25 (0.4) surface exact
# keyword/spec matches that embeddings alone might rank lower.
hybrid_retriever = EnsembleRetriever(
    retrievers=[bm25_retriever, vector_retriever],
    weights=[0.4, 0.6]
)

llm = ChatGoogleGenerativeAI(
    model="models/gemini-3.1-flash-lite",
    temperature=0.3,
    google_api_key=GOOGLE_API_KEY
)

class State(TypedDict):
    messages: Annotated[list[AnyMessage], add_messages]
    completed: bool
    user_role: str
    user_location: str
memory = MemorySaver()


prompt = ChatPromptTemplate.from_messages([
    (
        "system",
"""
You are a helpful medical equipment assistant.

User context:
- Role: {user_role}
- Location: {user_location}

Use this context to personalize your answers. For example:
- If the user is a doctor, use clinical language
- If the user is a buyer/procurement officer, focus on pricing and availability
- Mention location-specific shipping or availability if relevant

Answer ONLY using retrieved documents and chat history.
If the answer is not in the documents, say: "Sorry! I do not have that information."

At the end of every response, generate 3 short follow-up options.

Return STRICTLY in this JSON format:
{{
  "answer": "main chatbot answer",
  "options": ["option 1", "option 2", "option 3"]
}}

Retrieved Documents:
{retrieved_docs}
"""
    ),
    MessagesPlaceholder(variable_name="chat_history"),
    ("human", "{question}")
])
chain = prompt | llm  
# Update chain invoke in chat() to accept context
def chat(state: State):
    question = state["messages"][-1].content
    # --- CHANGED: was vector_store.similarity_search(question, k=7) ---
    retrieved_docs = hybrid_retriever.invoke(question)
    response = chain.invoke({
        "retrieved_docs": retrieved_docs,
        "question": question,
        "chat_history": state["messages"],
        "user_role": state.get("user_role", "general user"),
        "user_location": state.get("user_location", "unknown")
    })
    return {
        "messages": [AIMessage(content=response.content)],
        "completed": True
    }

graph_builder = StateGraph(State)
graph_builder.add_node("chat", chat)
graph_builder.add_edge(START, "chat")
graph_builder.add_edge("chat", END)
graph = graph_builder.compile(
    checkpointer=memory
)

def get_response(user_input, thread_id, user_role="general user", user_location="unknown"):
    config = {"configurable": {"thread_id": thread_id}}
    result = graph.invoke(
        {
            "messages": [HumanMessage(content=user_input)],
            "user_role": user_role,
            "user_location": user_location
        },
        config=config
    )
    raw_response = result["messages"][-1].content
    if isinstance(raw_response, list):
        raw_response = raw_response[0]["text"]
    try:
        parsed = json.loads(raw_response)
        return parsed.get("answer", ""), parsed.get("options", [])
    except Exception as e:
        print("JSON PARSE ERROR:", e)
        return raw_response, []