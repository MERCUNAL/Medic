import os
from pathlib import Path
import time

from dotenv import load_dotenv

from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

from docx import Document as DocxDocument


load_dotenv()


# =========================
# EMBEDDINGS
# =========================
embeddings = GoogleGenerativeAIEmbeddings(
    model="models/gemini-embedding-001"
)


# =========================
# CHROMADB
# =========================
vector_store = Chroma(
    collection_name="medical_data",
    embedding_function=embeddings,
    persist_directory="./chroma_db"
)


# =========================
# EXISTING IDS
# =========================
existing = vector_store.get()

existing_ids = set(existing["ids"])


# =========================
# FUNCTION TO EMBED DOCX
# =========================
def embed_docx(file_path, base_id):

    if not os.path.exists(file_path):

        print(f"{file_path} not found.")
        return

    doc = DocxDocument(file_path)

    full_text = []

    for para in doc.paragraphs:

        if para.text.strip():

            full_text.append(para.text.strip())

    # =========================
    # CHUNKING
    # =========================
    chunks = []

    current_chunk = ""

    for para in full_text:

        current_chunk += para + "\n"

        if len(current_chunk) > 500:

            chunks.append(current_chunk)
            current_chunk = ""

    if current_chunk:

        chunks.append(current_chunk)

    # =========================
    # ADD CHUNKS
    # =========================
    for i, chunk in enumerate(chunks):

        chunk_id = f"{base_id}_{i}"

        if chunk_id in existing_ids:

            print(f"{chunk_id} already exists.")
            continue

        print(f"Embedding {chunk_id}")

        document = Document(
            page_content=chunk,
            metadata={
                "source": file_path,
                "id": chunk_id
            }
        )

        vector_store.add_documents(
            documents=[document],
            ids=[chunk_id]
        )

        time.sleep(2)

    print(f"{file_path} embedded successfully.")


# =========================
# EMBED FILES
# =========================
from pathlib import Path
import pandas as pd
BASE_DIR = Path(__file__).resolve().parent.parent
    
embed_docx(
    BASE_DIR / "documents" / "Log and sign.docx",
    "log_and_sign"
)

embed_docx(
    BASE_DIR / "documents" / "FAQ.docx",
    "faq"
)