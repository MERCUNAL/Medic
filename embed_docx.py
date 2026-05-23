import os
import time
from dotenv import load_dotenv
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from docx import Document as DocxDocument
load_dotenv()

embeddings = GoogleGenerativeAIEmbeddings(
    model="models/gemini-embedding-001"
)

vector_store = Chroma(
    collection_name="medical_data",
    embedding_function=embeddings,
    persist_directory="./chroma_db"
)

existing = vector_store.get()
existing_ids = set(existing["ids"])


docx_file = "Log and sign.docx"
doc_id = "log_and_sign_doc"
faq_file = "FAQ.docx"
faq_id = "faq_doc"

if doc_id in existing_ids:
    print("Login and Signup Document already embedded.")

else:
    print("Embedding new document...")
    doc = DocxDocument(docx_file)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    docx_text = "\n".join(full_text)
    document = Document(
        page_content=docx_text,
        metadata={
            "source": "Log and sign.docx",
            "id": doc_id
        }
    )
if faq_id in existing_ids:
    print("FAQ Document already embedded.")

else:
    print("Embedding new document...")
    doc = DocxDocument(faq_file)
    full_text1 = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text1.append(para.text)
    docx_text = "\n".join(full_text1)
    document = Document(
        page_content=docx_text,
        metadata={
            "source": "FAQ.docx",
            "id": faq_id
        }
    )
    vector_store.add_documents(
        documents=[document],
        ids=[doc_id]
    )
    time.sleep(3)
    print("Document embedded successfully.")