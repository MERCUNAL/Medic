import pandas as pd
from langchain_chroma import Chroma
import json
from langchain_core.documents import Document
from docx import Document as DocxDocument
import os
from dotenv import load_dotenv
from langchain_google_genai import (
    GoogleGenerativeAIEmbeddings,
    ChatGoogleGenerativeAI
)
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_core.documents import Document
from langchain_core.prompts import (
    ChatPromptTemplate,
    MessagesPlaceholder
)
from langchain_core.messages import (
    HumanMessage,
    AIMessage,
    AnyMessage
)
from langgraph.graph.message import add_messages
from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import StateGraph, START, END
from typing_extensions import TypedDict, Annotated
load_dotenv()


def load_vector_store():

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001"
    )

    persist_directory = "./chroma_db"

    # LOAD CHROMADB
    vector_store = Chroma(
        collection_name="medical_data",
        embedding_function=embeddings,
        persist_directory=persist_directory
    )

    # =========================
    # GET EXISTING IDS
    # =========================
    existing = vector_store.get()

    existing_ids = set(existing["ids"])

    documents = []
    ids = []

    # =========================
    # LOAD EXCEL
    # =========================
    df = pd.read_excel("Medical_list.xlsx")

    for index, row in df.iterrows():

        row_dict = row.to_dict()

        chunk_text = json.dumps(row_dict)

        doc_id = f"excel_{index}"

        documents.append(
            Document(
                page_content=chunk_text,
                metadata={
                    "source": "excel",
                    "id": doc_id
                }
            )
        )

        ids.append(doc_id)

    # # =========================
    # # LOAD DOCX
    # # =========================
    # docx_file = "Log and sign.docx"

    # if os.path.exists(docx_file):

    #     doc = DocxDocument(docx_file)

    #     full_text = []

    #     for para in doc.paragraphs:

    #         if para.text.strip():

    #             full_text.append(para.text)

    #     docx_text = "\n".join(full_text)

    #     doc_id = "log_and_sign_doc"

    #     documents.append(
    #         Document(
    #             page_content=docx_text,
    #             metadata={
    #                 "source": "docx",
    #                 "id": doc_id
    #             }
    #         )
    #     )

    #     ids.append(doc_id)

    # # =========================
    # # FILTER NEW DOCS
    # # =========================
    # new_docs = []
    # new_ids = []

    # for doc, doc_id in zip(documents, ids):

    #     if doc_id not in existing_ids:

    #         new_docs.append(doc)
    #         new_ids.append(doc_id)

    # # =========================
    # # ADD ONLY NEW DOCS
    # # =========================
    # if new_docs:

    #     print(f"ADDING {len(new_docs)} NEW DOCUMENTS")

    #     vector_store.add_documents(
    #         documents=new_docs,
    #         ids=new_ids
    #     )

    # else:

    #     print("NO NEW DOCUMENTS TO EMBED")

    return vector_store
vector_store = load_vector_store()

llm = ChatGoogleGenerativeAI(
    model="models/gemini-3.1-flash-lite",
    temperature=0.3
)

class State(TypedDict):
    messages: Annotated[list[AnyMessage], add_messages]
    completed: bool
memory = MemorySaver()


prompt = ChatPromptTemplate.from_messages([
    (
        "system",
        """
        You are a medical chatbot.
        Communicate with the user like a chatbot, providing helpful and accurate information.

        Answer the user's question using only:
        1. Retrieved documents
        2. Chat history

        Retrieved Documents:
        {retrieved_docs}
        """
    ),
    MessagesPlaceholder(variable_name="chat_history"),

    (
        "human",
        "{question}"
    )
])
chain = prompt | llm

def chat(state: State):
    question = state["messages"][-1].content
    retrieved_docs = vector_store.similarity_search(
        question,
        k=3
    )
    response = chain.invoke({
        "retrieved_docs": retrieved_docs,
        "question": question,
        "chat_history": state["messages"]
    })
    return {
        "messages": [AIMessage(content=response.content)],
        "completed": True
    }


graph_builder = StateGraph(State)
graph_builder.add_node("chat", chat)
graph_builder.add_edge(START, "chat")
graph_builder.add_edge("chat", END)
graph = graph_builder.compile(
    checkpointer=memory
)

def get_response(user_input, thread_id):
    config = {
        "configurable": {
            "thread_id": thread_id
        }
    }
    result = graph.invoke(
        {
            "messages": [
                HumanMessage(content=user_input)
            ]
        },
        config=config
    )
    return result["messages"][-1].content[0]['text']




